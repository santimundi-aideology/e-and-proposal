"""
Build a TWO-PAGE Executive Briefing (Word .docx) for e& management.

Scope:
    - 2-pillar engagement (SMB + Enterprise & Government) at a glance
    - All 6 SMB agents and 3 Enterprise tiers, with first GTM cost proposal
    - Commercial headlines (no AIdeology financials, no break-even)

Output:
    Internal Analysis/eand_Executive_Briefing_AIdeology_Engagement_v0.1.docx
"""
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


OUT = "Internal Analysis/eand_Executive_Briefing_AIdeology_Engagement_v0.1.docx"

# Brand palette
RED         = RGBColor(0xE0, 0x08, 0x00)
BLACK       = RGBColor(0x11, 0x11, 0x11)
INK         = RGBColor(0x22, 0x22, 0x22)
GREY        = RGBColor(0x66, 0x66, 0x66)
HEADER_FILL = "F5F5F5"
BAND_FILL   = "FAFAFA"


def set_cell_shading(cell, hex_color):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tc_pr.append(shd)


def set_cell_borders(cell, color="DDDDDD", size="4"):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_borders = OxmlElement("w:tcBorders")
    for edge in ("top", "left", "bottom", "right"):
        b = OxmlElement(f"w:{edge}")
        b.set(qn("w:val"), "single")
        b.set(qn("w:sz"), size)
        b.set(qn("w:color"), color)
        tc_borders.append(b)
    tc_pr.append(tc_borders)


def add_paragraph_border(paragraph, position="bottom", color="E00800", size="14"):
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = OxmlElement("w:pBdr")
    border = OxmlElement(f"w:{position}")
    border.set(qn("w:val"), "single")
    border.set(qn("w:sz"), size)
    border.set(qn("w:space"), "3")
    border.set(qn("w:color"), color)
    p_bdr.append(border)
    p_pr.append(p_bdr)


def style_run(run, *, size=10, bold=False, italic=False, color=INK, font="Calibri"):
    run.font.name = font
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    run.font.color.rgb = color


def p(doc, text, *, size=9.5, bold=False, italic=False, color=INK,
      align=WD_ALIGN_PARAGRAPH.LEFT, space_before=0, space_after=2):
    para = doc.add_paragraph()
    para.alignment = align
    para.paragraph_format.space_before = Pt(space_before)
    para.paragraph_format.space_after = Pt(space_after)
    para.paragraph_format.line_spacing = 1.15
    run = para.add_run(text)
    style_run(run, size=size, bold=bold, italic=italic, color=color)
    return para


def section(doc, eyebrow, title):
    """Compact section header — red eyebrow + bold title with bottom rule."""
    h = doc.add_paragraph()
    h.paragraph_format.space_before = Pt(8)
    h.paragraph_format.space_after = Pt(4)
    h.paragraph_format.line_spacing = 1.0
    r = h.add_run(f"{eyebrow.upper()}    ")
    style_run(r, size=8, bold=True, color=RED)
    r = h.add_run(title)
    style_run(r, size=11.5, bold=True, color=BLACK)
    add_paragraph_border(h, "bottom", color="E00800", size="10")


def make_table(doc, headers, rows, *, widths=None, first_col_bold=True,
               row_size=8.8, header_size=7.8):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False

    if widths:
        for i, w in enumerate(widths):
            for cell in table.columns[i].cells:
                cell.width = Cm(w)

    # Header
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        set_cell_shading(cell, HEADER_FILL)
        set_cell_borders(cell, color="CCCCCC", size="4")
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        para = cell.paragraphs[0]
        para.paragraph_format.space_before = Pt(1)
        para.paragraph_format.space_after = Pt(1)
        para.paragraph_format.line_spacing = 1.0
        run = para.add_run(h.upper())
        style_run(run, size=header_size, bold=True, color=GREY)

    # Body
    for r_idx, row in enumerate(rows):
        for c_idx, val in enumerate(row):
            cell = table.rows[r_idx + 1].cells[c_idx]
            if r_idx % 2 == 1:
                set_cell_shading(cell, BAND_FILL)
            set_cell_borders(cell, color="EEEEEE", size="4")
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            para = cell.paragraphs[0]
            para.paragraph_format.space_before = Pt(1)
            para.paragraph_format.space_after = Pt(1)
            para.paragraph_format.line_spacing = 1.1
            run = para.add_run(str(val))
            is_first_col = c_idx == 0
            style_run(
                run,
                size=row_size,
                bold=first_col_bold and is_first_col,
                color=BLACK if is_first_col else INK,
            )


# ============================================================================
# Build
# ============================================================================
doc = Document()

style = doc.styles["Normal"]
style.font.name = "Calibri"
style.font.size = Pt(9.5)

for s in doc.sections:
    s.top_margin    = Cm(1.3)
    s.bottom_margin = Cm(1.3)
    s.left_margin   = Cm(1.5)
    s.right_margin  = Cm(1.5)


# ---- Title block (compact) ----
eyebrow = doc.add_paragraph()
eyebrow.paragraph_format.space_after = Pt(0)
r = eyebrow.add_run("EXECUTIVE BRIEFING  ·  FOR e& MANAGEMENT  ·  CONFIDENTIAL")
style_run(r, size=8, bold=True, color=RED)

title = doc.add_paragraph()
title.paragraph_format.space_before = Pt(2)
title.paragraph_format.space_after = Pt(2)
r = title.add_run("e& + AIdeology — Agentic AI Engagement")
style_run(r, size=18, bold=True, color=BLACK)
add_paragraph_border(title, "bottom", color="E00800", size="14")

sub = doc.add_paragraph()
sub.paragraph_format.space_before = Pt(2)
sub.paragraph_format.space_after = Pt(0)
r = sub.add_run(
    "A two-pillar partnership to make e& the destination for AI across the UAE and the OpCo footprint: "
    "the SMB Digital Marketplace (six pre-built agents on the e& invoice) and the Enterprise & Government "
    "AI delivery catalogue (three tiers, from adapted agents to sovereign on-prem)."
)
style_run(r, size=9.5, italic=True, color=GREY)


# ---- 1. Engagement at a glance ----
section(doc, "01", "Engagement at a glance")
make_table(
    doc,
    headers=["", "Pillar 01 — SMB Digital Marketplace", "Pillar 02 — Enterprise & Government"],
    rows=[
        ["Buyer",       "320K+ paying SMBs already on the e& monthly invoice.",                                 "UAE banks, hospitals, regulators, energy majors, federal entities."],
        ["Offer",       "6 pre-built agents on a subscription marketplace (Spark / Scale / Command tiers).",     "3 tiers — adapted SMB agents (T1), custom multi-agent solutions (T2), sovereign on-prem (T3)."],
        ["Motion",       "Self-service marketplace + e& SMB sales force. Multi-tenant cloud.",                   "e& Enterprise team fronts the customer; AIdeology designs and delivers the AI plane."],
        ["GTM model",   "$3.44M one-off platform + 6-agent build (Wave 1–5, 30 weeks).",                          "Per-customer build fee + monthly managed service. Hosting & connectivity 100% e&."],
    ],
    widths=[2.3, 8.0, 8.0],
    row_size=8.8,
)


# ---- 2. SMB pillar — six agents & GTM cost ----
section(doc, "02", "Pillar 01 — Six SMB agents, subscription pricing, and GTM build cost")
make_table(
    doc,
    headers=["Agent", "What it does", "Subscription (AED/mo)", "GTM build cost", "Wave"],
    rows=[
        ["Customer", "Voice + WhatsApp + web in Arabic & English. Books, answers, escalates 24/7 on Toll Free 800 + e& BSP.",            "199 / 349 / 999", "$1,250,000*", "Wave 1"],
        ["Sales",    "Native CRM + lead scoring + AI follow-ups + proposal drafting. Syncs to Dynamics, Salesforce, HubSpot, Zoho.",     "149 / 349 / 699", "$350,000",    "Wave 2"],
        ["Comms",    "Unified inbox + AI campaign builder across WhatsApp, SMS, email, Teams. e& Smart Messaging + BSP carrier-grade.",   "99 / 299 / 649",  "$350,000",    "Wave 2"],
        ["Finance",  "Invoices, e& Pay links, FTA VAT, AI reminders, cash-flow forecast. Syncs to QuickBooks, Xero, Business Central.",    "199 / 449 / 849", "$300,000",    "Wave 3"],
        ["Ops",      "Tasks, approvals, SOP knowledge base, daily WhatsApp summaries. Syncs to Teams, SharePoint, Google Workspace.",      "149 / 299 / 599", "$300,000",    "Wave 3"],
        ["People",   "WPS payroll, attendance via e& SIM, leave, onboarding, visa-expiry alerts, AI CV screening. Connects MOHRE/GDRFA.",  "199 / 449 / 849", "$275,000",    "Wave 4"],
        ["Platform foundation, hardening & 3-year L3 support", "Wave 1 platform foundation + Wave 5 security audit, pen test, runbooks, documentation, formal handoff, 3-year L3 platform support.", "—", "$618,621",  "Wave 1 + 5"],
        ["TOTAL — all 6 agents + platform + 3-year support",   "Full SMB programme delivered in 30 weeks, all six agents in production, platform handed over to e&.",                              "—", "$3,443,621", "All waves"],
    ],
    widths=[2.5, 9.8, 2.5, 2.4, 1.5],
    row_size=8.5,
)
p(doc,
  "* The Customer Agent is bundled inside Wave 1 with the orchestration platform foundation — the first agent absorbs the platform investment that every later agent reuses. Wave-5 hardening and 3-year L3 platform support are also included in the totals.",
  size=8, italic=True, color=GREY, space_after=2)


# ---- 3. Enterprise pillar — three tiers & GTM cost ----
section(doc, "03", "Pillar 02 — Three enterprise tiers, sold per customer")
make_table(
    doc,
    headers=["Tier", "What e& sells", "GTM cost — per customer", "Managed service"],
    rows=[
        ["T1 — Adapted Agents",
         "The six SMB agents, enterprise-hardened for one customer: multi-branch routing, deep ERP/CRM/HIS connectors, RBAC, audit, domain fine-tuning. First agent live in 60 days.",
         "Agent #1: $60–80K  ·  #2: $36–48K (−40%)  ·  #3+: $30–40K (−50%)",
         "$5–15K / month"],
        ["T2 — Custom AI",
         "Multi-agent solutions from spec. Headline: AI Contact Centre. Catalogue also: Document Intelligence, Approval Orchestration, Predictive Maintenance, Security & Compliance.",
         "$300K–$680K depending on family (AI Contact Centre $550–680K)",
         "$12–45K / month"],
        ["T3 — Sovereign On-Prem",
         "Tier 2 deployed on the customer's site — air-gap capable, NVIDIA-Certified hardware, Help AG security wrap, customer-owned keys.",
         "AIdeology build $600K–$1M  ·  Hardware at cost (HPC RA library)  ·  Help AG $80–220K",
         "$25–60K / month"],
    ],
    widths=[3.5, 7.5, 5.5, 2.5],
    row_size=8.6,
)
p(doc,
  "Hosting, connectivity and Tier-3 hardware are 100% e& — full margin retained by e& on top of the build and managed-service fees.",
  size=8.2, italic=True, color=GREY, space_after=2)


# ---- 4. Commercial headlines ----
section(doc, "04", "Commercial headlines — what e& gets")
make_table(
    doc,
    headers=["Topic", "Headline"],
    rows=[
        ["SMB revenue share",     "65 / 35 in Year 1–2  →  72 / 28 in Year 3  →  80 / 20 in Year 4+ (e& majority from Day 1; declining as e& team ramps)."],
        ["Enterprise revenue split","60 / 40 AIdeology / e& on build and managed service. Hosting, connectivity and hardware 100% e&."],
        ["IP ownership",          "Agent IP transfers to e& by end of Year 2. Platform IP stays with AIdeology (perpetual non-exclusive licence to e&; buyout option from Year 3)."],
        ["Support tiers",          "L1 + L2 owned by e& from Day 1. L3 + L4 (platform engineering) owned by AIdeology — included in the fixed fee for Years 1–3."],
        ["Build-then-transfer",    "AIdeology builds and trains; e& progressively takes ownership. 2–3 e& engineers embedded Year 1, 6–8 by Year 3, full ownership by Year 4."],
        ["OpCo expansion",         "After UAE proof, e& OpCo teams localise and deploy themselves — no second build fee per country. Saudi → Morocco → Egypt / Kuwait → rest."],
    ],
    widths=[3.5, 15.0],
    row_size=8.7,
)


# ---- 5. Next steps ----
section(doc, "05", "Next steps")
make_table(
    doc,
    headers=["#", "Action", "Owner", "Timeline"],
    rows=[
        ["1", "Endorsement of the two-pillar engagement by e& Management — formal sponsorship of SMB + Enterprise programmes.", "e& EXCO sponsor",               "Within 2 weeks"],
        ["2", "Working session on Wave 1 SDD, infrastructure (G42 / Azure), and the embedded e& engineering team.",              "AIdeology + e& AI / engineering","Within 4 weeks"],
        ["3", "Identification of the first 3 enterprise / government anchor accounts for Tier 1 / Tier 2 / Tier 3 deployments.", "e& Enterprise account team",     "Within 4 weeks"],
    ],
    widths=[0.8, 11.5, 3.7, 2.5],
    row_size=8.6,
)

p(doc,
  "On signature, AIdeology mobilises the Wave 1 team within 14 days and delivers the platform MVP and Customer Agent in production inside the e& environment within 12 weeks.",
  size=8.5, italic=True, color=GREY, space_before=4, space_after=0)


doc.save(OUT)
print(f"Saved: {OUT}")
